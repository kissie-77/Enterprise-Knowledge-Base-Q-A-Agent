"""
RAG 增强检索生成引擎
参考 hello-agents 的 RAGTool 设计，实现完整的文档处理与检索管道

特性：
- 多格式文档解析（PDF/TXT/DOCX/MD）
- 智能文本分块（语义感知）
- 向量化存储（ChromaDB）
- 多查询扩展（MQE）
- 相关性评分与重排序
"""
import os
import re
import uuid
from typing import List, Dict, Any, Optional

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter


# 默认配置
DEFAULT_CHROMA_PATH = "chroma_store"
DEFAULT_COLLECTION = "knowledge_base"
DEFAULT_EMBEDDING_MODEL = "all-MiniLM-L6-v2"


class RAGEngine:
    """
    RAG 增强检索生成引擎

    使用方式:
        rag = RAGEngine()

        # 添加文档
        await rag.add_document("path/to/file.pdf")

        # 检索相关内容
        results = await rag.search("用户问题", top_k=5)

        # 构建上下文
        context = rag.build_context(results)
    """

    def __init__(
        self,
        chroma_path: str = DEFAULT_CHROMA_PATH,
        collection_name: str = DEFAULT_COLLECTION,
        embedding_model_name: str = DEFAULT_EMBEDDING_MODEL,
    ):
        """
        初始化 RAG 引擎

        Args:
            chroma_path: ChromaDB 持久化路径
            collection_name: 集合名称
            embedding_model_name: 嵌入模型名称
        """
        self.chroma_path = chroma_path
        self.collection_name = collection_name
        self._embedding_model_name = embedding_model_name
        self._embedding_model: Optional[SentenceTransformer] = None

        # 确保存储目录存在
        os.makedirs(chroma_path, exist_ok=True)

        print(f"📚 RAG 引擎初始化: collection={collection_name}, embedding={embedding_model_name}")

    @property
    def embedding_model(self) -> SentenceTransformer:
        """懒加载嵌入模型"""
        if self._embedding_model is None:
            print(f"⏳ 正在加载嵌入模型: {self._embedding_model_name}...")
            self._embedding_model = SentenceTransformer(self._embedding_model_name)
            print("✅ 嵌入模型加载完成")
        return self._embedding_model

    @property
    def chroma_client(self) -> chromadb.ClientAPI:
        """获取 ChromaDB 客户端"""
        return chromadb.PersistentClient(
            path=self.chroma_path,
            settings=Settings(anonymized_telemetry=False, allow_reset=True),
        )

    @property
    def collection(self):
        """获取当前集合"""
        return self.chroma_client.get_or_create_collection(name=self.collection_name)

    # ===== 文档处理 =====

    async def add_document(
        self,
        file_path: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> int:
        """
        处理文档并添加到向量库

        Args:
            file_path: 文件路径
            chunk_size: 分块大小
            chunk_overlap: 分块重叠
            metadata: 额外元数据

        Returns:
            添加的文档片段数
        """
        # 1. 提取文本
        text = self._extract_text(file_path)
        if not text:
            raise ValueError(f"无法从文件提取内容: {file_path}")

        # 2. 清理文本
        text = self._clean_text(text)

        # 3. 智能分块
        chunks = self._split_text(text, chunk_size, chunk_overlap)
        if not chunks:
            raise ValueError("文本分块失败")

        # 4. 向量化并存储
        await self._store_chunks(chunks, file_path, metadata)

        print(f"✅ 文档处理完成: {os.path.basename(file_path)} → {len(chunks)} 个片段")
        return len(chunks)

    async def add_text(
        self,
        text: str,
        document_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ) -> int:
        """
        直接添加文本到向量库

        Args:
            text: 文本内容
            document_id: 文档标识
            metadata: 元数据
            chunk_size: 分块大小
            chunk_overlap: 分块重叠

        Returns:
            添加的片段数
        """
        text = self._clean_text(text)
        chunks = self._split_text(text, chunk_size, chunk_overlap)

        if not chunks:
            return 0

        await self._store_chunks(chunks, document_id or "direct_text", metadata)
        return len(chunks)

    # ===== 检索 =====

    async def search(
        self,
        query: str,
        top_k: int = 5,
        score_threshold: Optional[float] = None,
    ) -> List[Dict[str, Any]]:
        """
        检索与查询相关的文档片段

        Args:
            query: 查询文本
            top_k: 返回结果数量
            score_threshold: 相关性分数阈值（越小越相似，基于 L2 距离）

        Returns:
            检索结果列表，每项包含 content, score, metadata
        """
        # 向量化查询
        query_embedding = self.embedding_model.encode([query])[0].tolist()

        # 检索
        collection = self.collection
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
        )

        # 解析结果
        chunks = []
        if results["documents"] and results["documents"][0]:
            for i, doc in enumerate(results["documents"][0]):
                score = results["distances"][0][i] if results.get("distances") else 0
                meta = results["metadatas"][0][i] if results.get("metadatas") else {}

                # 分数过滤
                if score_threshold and score > score_threshold:
                    continue

                chunks.append({
                    "content": doc,
                    "score": score,
                    "metadata": meta,
                })

        return chunks

    async def search_with_mqe(
        self,
        query: str,
        top_k: int = 5,
        num_expansions: int = 3,
    ) -> List[Dict[str, Any]]:
        """
        多查询扩展检索（Multi-Query Expansion）
        通过多个相关查询扩大检索范围，提高召回率

        Args:
            query: 原始查询
            top_k: 每个查询返回的结果数
            num_expansions: 扩展查询数量

        Returns:
            去重后的检索结果
        """
        # 生成扩展查询
        expanded_queries = self._expand_query(query, num_expansions)
        all_queries = [query] + expanded_queries

        # 合并检索结果（去重）
        seen_contents = set()
        all_results = []

        for q in all_queries:
            results = await self.search(q, top_k=top_k)
            for result in results:
                content_hash = hash(result["content"][:100])
                if content_hash not in seen_contents:
                    seen_contents.add(content_hash)
                    all_results.append(result)

        # 按分数排序（L2 距离越小越相似）
        all_results.sort(key=lambda x: x["score"])
        return all_results[:top_k * 2]

    # ===== 上下文构建 =====

    def build_context(
        self,
        results: List[Dict[str, Any]],
        max_length: int = 3000,
    ) -> str:
        """
        从检索结果构建上下文文本

        Args:
            results: 检索结果列表
            max_length: 上下文最大长度

        Returns:
            格式化的上下文字符串
        """
        if not results:
            return "知识库中未找到相关内容。"

        context_parts = []
        current_length = 0

        for i, result in enumerate(results):
            content = result["content"]
            # 截断过长的内容
            if current_length + len(content) > max_length:
                remaining = max_length - current_length
                if remaining > 100:
                    content = content[:remaining] + "..."
                else:
                    break

            context_parts.append(f"[参考资料{i+1}] {content}")
            current_length += len(content)

        return "\n\n".join(context_parts)

    # ===== 管理 =====

    def get_stats(self) -> Dict[str, Any]:
        """获取向量库统计信息"""
        collection = self.collection
        count = collection.count()
        return {
            "collection_name": self.collection_name,
            "total_chunks": count,
            "embedding_model": self._embedding_model_name,
            "chroma_path": self.chroma_path,
        }

    def reset(self):
        """重置向量库（清空所有数据）"""
        client = self.chroma_client
        try:
            client.delete_collection(self.collection_name)
            print(f"🗑️ 集合 '{self.collection_name}' 已删除")
        except Exception:
            pass

    # ===== 内部方法 =====

    def _extract_text(self, file_path: str) -> str:
        """从文件提取文本"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext == ".txt" or ext == ".md":
            return self._extract_txt(file_path)
        elif ext == ".docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}。支持: PDF, TXT, MD, DOCX")

    def _extract_pdf(self, file_path: str) -> str:
        """提取 PDF 文本"""
        try:
            import PyPDF2
            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            print(f"❌ PDF 提取失败: {e}")
            return ""

    def _extract_txt(self, file_path: str) -> str:
        """提取文本文件内容"""
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        return ""

    def _extract_docx(self, file_path: str) -> str:
        """提取 Word 文档文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"❌ DOCX 提取失败: {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除特殊控制字符
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
        # 合并多余空白
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    def _split_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """智能文本分块"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n",   # 段落
                "\n",     # 换行
                "。",     # 中文句号
                "！",     # 中文感叹号
                "？",     # 中文问号
                ".",      # 英文句号
                "!",      # 英文感叹号
                "?",      # 英文问号
                "；",     # 中文分号
                ";",      # 英文分号
                "，",     # 中文逗号
                ",",      # 英文逗号
                " ",      # 空格
            ],
        )
        return splitter.split_text(text)

    async def _store_chunks(
        self,
        chunks: List[str],
        source: str,
        metadata: Optional[Dict[str, Any]] = None,
    ):
        """存储文本块到向量库"""
        # 生成嵌入向量
        embeddings = self.embedding_model.encode(chunks).tolist()

        # 准备 ID 和元数据
        ids = [str(uuid.uuid4()) for _ in range(len(chunks))]
        metadatas = []
        for i, chunk in enumerate(chunks):
            meta = {
                "source": os.path.basename(source) if os.path.exists(source) else source,
                "chunk_index": i,
                "char_count": len(chunk),
            }
            if metadata:
                meta.update({k: str(v) for k, v in metadata.items()})
            metadatas.append(meta)

        # 存入 ChromaDB
        collection = self.collection
        collection.add(
            documents=chunks,
            embeddings=embeddings,
            ids=ids,
            metadatas=metadatas,
        )

    def _expand_query(self, query: str, num_expansions: int = 3) -> List[str]:
        """
        查询扩展（简单实现，基于关键词重组）
        更高级的实现可以用 LLM 生成扩展查询
        """
        expanded = []

        # 策略1: 去除停用词后的核心查询
        stopwords = {"的", "了", "是", "在", "和", "有", "与", "对", "中", "为", "什么", "如何", "怎么", "请", "吗"}
        words = [w for w in query if w not in stopwords]
        if words:
            expanded.append("".join(words))

        # 策略2: 添加同义表述
        if "是什么" in query:
            expanded.append(query.replace("是什么", "的定义和概念"))
        elif "如何" in query:
            expanded.append(query.replace("如何", "方法步骤"))
        elif "为什么" in query:
            expanded.append(query.replace("为什么", "原因分析"))

        # 策略3: 关键短语提取（简单实现）
        if len(query) > 10:
            mid = len(query) // 2
            expanded.append(query[:mid])
            expanded.append(query[mid:])

        return expanded[:num_expansions]
