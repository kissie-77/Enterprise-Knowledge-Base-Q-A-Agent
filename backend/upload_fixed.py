"""
文档上传与向量化处理模块
"""
import os
import re
from typing import List

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter
import PyPDF2
from docx import Document
import asyncio

# 配置
CHROMA_PATH = "chroma_store"
COLLECTION_NAME = "knowledge_base"

# 初始化嵌入模型（本地运行）
embedding_model = None


def get_embedding_model():
    """获取嵌入模型实例（懒加载）"""
    global embedding_model
    if embedding_model is None:
        print("正在加载嵌入模型...")
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        print("✅ 嵌入模型加载完成")
    return embedding_model


def get_chroma_client():
    """获取 Chroma 客户端"""
    client = chromadb.PersistentClient(
        path=CHROMA_PATH,
        settings=Settings(
            anonymized_telemetry=False,
            allow_reset=True
        )
    )
    return client


def extract_text_from_pdf(file_path: str) -> str:
    """从 PDF 文件提取文本"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"PDF 提取错误: {e}")
    return text


def extract_text_from_txt(file_path: str) -> str:
    """从 TXT 文件提取文本"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"TXT 读取错误: {e}")
        # 尝试其他编码
        try:
            with open(file_path, 'r', encoding='gbk') as file:
                return file.read()
        except:
            return ""


def extract_text_from_docx(file_path: str) -> str:
    """从 Word 文档提取文本"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # 也提取表格中的内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        return text
    except Exception as e:
        print(f"Word 文档读取错误: {e}")
        return ""


def clean_text(text: str) -> str:
    """清理文本，移除多余空白字符"""
    # 移除多余的空白字符
    text = re.sub(r'\s+', ' ', text)
    # 移除特殊字符
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text.strip()


def split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """文本切片"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        length_function=len,
        separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", "；", ";", "，", ","]
    )
    chunks = text_splitter.split_text(text)
    return chunks


async def process_document(file_path: str, chunk_size: int = 500, overlap: int = 50) -> int:
    """
    处理文档：提取文本、切片、向量化并存入 Chroma

    Args:
        file_path: 文件路径
        chunk_size: 切片大小
        overlap: 重叠字符数

    Returns:
        切片数量
    """
    # 提取文本
    if file_path.lower().endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.txt'):
        text = extract_text_from_txt(file_path)
    elif file_path.lower().endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("不支持的文件类型，仅支持 PDF、TXT、DOCX 格式")

    if not text:
        raise ValueError("无法提取文档内容")

    # 清理文本
    text = clean_text(text)

    # 切片
    chunks = split_text(text, chunk_size, overlap)

    if not chunks:
        raise ValueError("文档切片失败")

    # 获取嵌入模型
    model = get_embedding_model()

    # 生成向量
    embeddings = model.encode(chunks).tolist()

    # 存入 Chroma
    client = get_chroma_client()
    collection = client.get_or_create_collection(name=COLLECTION_NAME)

    # 生成唯一 ID
    import uuid
    ids = [str(uuid.uuid4()) for _ in range(len(chunks))]

    # 添加到集合
    collection.add(
        documents=chunks,
        embeddings=embeddings,
        ids=ids
    )

    print(f"✅ 文档处理完成：{len(chunks)} 个片段已存入向量库")

    return len(chunks)
